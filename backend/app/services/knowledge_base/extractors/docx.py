import io

from docx import Document as DocxDocument

from app.core.exceptions import ExtractionError
from app.services.knowledge_base.extractors.base import ExtractedSegment


def extract_docx(content: bytes, filename: str) -> list[ExtractedSegment]:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionError(filename, str(exc)) from exc

    segments: list[ExtractedSegment] = []
    current_section: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        text = "\n".join(buffer).strip()
        if text:
            segments.append(ExtractedSegment(text=text, page=None, section_title=current_section))
        buffer.clear()

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            flush()
            current_section = text
        else:
            buffer.append(text)
    flush()

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            segments.append(ExtractedSegment(text="\n".join(rows), page=None, section_title=current_section))

    return segments
